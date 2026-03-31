#!/usr/bin/env python3
"""
Genera evaluación EYR en el formato visual del sistema (azul, celdas oscuras).
Uso: python3 scripts/generar_plantilla_evaluacion.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── colores ──────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1F, 0x38, 0x64)   # azul oscuro marino (títulos/secciones)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
FONT   = 'Arial'


# ─── helpers ──────────────────────────────────────────────────────────────────

def run(para, text, bold=False, italic=False, size=10, color=None, underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    r.font.name  = FONT
    r.font.size  = Pt(size)
    if color:
        r.font.color.rgb = color
    return r

def nospace(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def set_cell_bg(cell, hex_color: str):
    """Fondo de celda: hex_color como '1F3864' o '000000'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders_all(cell, color='000000', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)

def add_table_borders(tbl, sz=8):
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        b.append(el)
    tblPr.append(b)

def replace_tblgrid(tbl, twip_widths):
    tbl_xml = tbl._tbl
    for old in tbl_xml.findall(qn('w:tblGrid')):
        tbl_xml.remove(old)
    grid = OxmlElement('w:tblGrid')
    for w in twip_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        grid.append(gc)
    tbl_xml.insert(2, grid)

def set_cell_widths(tbl, twip_widths):
    replace_tblgrid(tbl, twip_widths)
    for row in tbl.rows:
        seen = set()
        for i, cell in enumerate(row.cells):
            if i < len(twip_widths) and id(cell) not in seen:
                cell.width = Emu(twip_widths[i] * 635)
                seen.add(id(cell))

def add_vmerge(cell, restart=False):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:vMerge')):
        tcPr.remove(old)
    vm = OxmlElement('w:vMerge')
    if restart:
        vm.set(qn('w:val'), 'restart')
    tcPr.append(vm)


# ─── secciones del documento ──────────────────────────────────────────────────

def add_school_header(doc, asignatura, curso, profesor):
    """Encabezado centrado: nombre colegio, ciudad, título, asignatura, profesor."""
    # Nombre colegio
    p = doc.add_paragraph()
    nospace(p, 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "CENTRO EDUCACIONAL ERNESTO YÁÑEZ RIVERA",
        bold=True, size=14, color=BLUE)

    # Ciudad
    p2 = doc.add_paragraph()
    nospace(p2, 0, 6)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, "Huechuraba · Santiago", size=9)


def add_eval_title(doc, titulo_eval, asignatura, curso, profesor):
    """Título de la evaluación + datos de asignatura/curso/profesor."""
    p = doc.add_paragraph()
    nospace(p, 0, 4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, titulo_eval, bold=True, size=18, color=BLUE)

    p2 = doc.add_paragraph()
    nospace(p2, 0, 2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, f"{asignatura}  |  {curso}", italic=True, size=10)

    p3 = doc.add_paragraph()
    nospace(p3, 0, 8)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p3, f"Profesor(a): {profesor}", size=10)


def add_info_table(doc, curso, anio, exigencia, puntaje_ideal):
    """
    Tabla 2 filas × 4 cols:
      F0: [Nombre: oscuro] [vacío] [Fecha: oscuro] [__/__/año]
      F1: [Curso: oscuro]  [valor] [Puntaje: oscuro] [___/pts]
    Celdas de etiqueta: fondo negro, texto blanco bold.
    """
    # anchos en twips: label≈1600, valor≈4000, label≈1600, valor≈2280 total≈9480
    COLS = [1600, 4000, 1800, 2080]

    tbl = doc.add_table(rows=2, cols=4)
    add_table_borders(tbl, sz=6)

    def label_cell(cell, text):
        set_cell_bg(cell, '1F3864')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        nospace(p, 2, 2)
        run(p, text, bold=True, size=10, color=WHITE)

    def value_cell(cell, text, bold=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        nospace(p, 2, 2)
        p.paragraph_format.left_indent = Pt(4)
        run(p, text, bold=bold, size=10)

    # Fila 0
    label_cell(tbl.cell(0, 0), "Nombre:")
    value_cell(tbl.cell(0, 1), "")
    label_cell(tbl.cell(0, 2), "Fecha:")
    value_cell(tbl.cell(0, 3), f"____  /  ____  / {anio}")

    # Fila 1
    label_cell(tbl.cell(1, 0), "Curso:")
    value_cell(tbl.cell(1, 1), curso)
    label_cell(tbl.cell(1, 2), "Puntaje:")
    value_cell(tbl.cell(1, 3), f"_______ / {puntaje_ideal}")

    set_cell_widths(tbl, COLS)
    doc.add_paragraph()


def add_oa_line(doc, oa_codigo, oa_texto):
    """OA como línea con etiqueta bold + texto."""
    p = doc.add_paragraph()
    nospace(p, 0, 4)
    run(p, "Objetivo de Aprendizaje:  ", bold=True, size=10)
    run(p, f"{oa_codigo} {oa_texto}", italic=True, size=10)


def add_general_instructions(doc, exigencia):
    p = doc.add_paragraph()
    nospace(p, 0, 8)
    run(p, f"Exigencia: {exigencia}   ·   Lea atentamente cada pregunta antes de contestar.",
        italic=True, size=10)


def add_section_header(doc, romano, titulo, instruccion_seccion=None):
    """
    Cabecera de sección azul bold (instrucción por ítem en negrita).
    Sub-instrucción opcional en itálica.
    """
    p = doc.add_paragraph()
    nospace(p, 10, 2)
    run(p, f"{romano}  {titulo.upper()}", bold=True, size=11, color=BLUE)

    if instruccion_seccion:
        p2 = doc.add_paragraph()
        nospace(p2, 0, 4)
        run(p2, instruccion_seccion, italic=True, size=10)


def add_question(doc, n, enunciado, opciones):
    p = doc.add_paragraph()
    nospace(p, 5, 2)
    p.paragraph_format.left_indent       = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run(p, f"{n}.  ", bold=True, size=10)
    run(p, enunciado, size=10)

    for letter, opcion in zip(['a)', 'b)', 'c)', 'd)'], opciones[:4]):
        op = doc.add_paragraph()
        nospace(op, 0, 1)
        op.paragraph_format.left_indent = Cm(1.4)
        run(op, f"{letter}  {opcion}", size=10)


def add_footer_line(doc):
    p = doc.add_paragraph()
    nospace(p, 20, 4)
    p.paragraph_format.left_indent = Cm(0)
    # línea separadora via borde inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:bottom')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '1F3864')
    pBdr.append(top)
    pPr.append(pBdr)

    p2 = doc.add_paragraph()
    nospace(p2, 2, 0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, "EYR Digital  ·  Centro Educacional Ernesto Yáñez Rivera  ·  Huechuraba",
        italic=True, size=8)


# ─── pauta ────────────────────────────────────────────────────────────────────

def add_pauta_title(doc, n_eval, titulo, curso):
    p = doc.add_paragraph()
    nospace(p, 8, 4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, f"PAUTA DE ESPECIFICACIÓN — EVALUACIÓN N°{n_eval}",
        bold=True, size=13, color=BLUE)
    p2 = doc.add_paragraph()
    nospace(p2, 0, 8)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p2, f"{titulo.upper()}  ·  {curso.upper()}",
        bold=True, size=10, color=BLUE)


def add_pauta_oa_box(doc, oa_codigo, oa_texto):
    tbl = doc.add_table(rows=1, cols=1)
    add_table_borders(tbl)
    set_cell_bg(tbl.cell(0,0), 'EBF0FA')
    p = tbl.cell(0,0).paragraphs[0]
    nospace(p, 4, 4)
    run(p, f"{oa_codigo}  ", bold=True, size=10)
    run(p, oa_texto, size=10)
    doc.add_paragraph()


def add_pauta_table(doc, secciones):
    total_q = sum(len(s['preguntas']) for s in secciones)
    # anchos twips: Indicador=3118, Habilidad=1587, Pregunta=1304, Clave=1134, Puntaje=1827
    COLS = [3118, 1587, 1304, 1134, 1827]

    tbl = doc.add_table(rows=1 + total_q + 1, cols=5)
    add_table_borders(tbl)

    def hdr_cell(cell, text):
        set_cell_bg(cell, '1F3864')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        nospace(p, 3, 3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, text, bold=True, size=9, color=WHITE)

    # Encabezado
    for i, h in enumerate(['Indicadores', 'Habilidad', 'N°', 'Clave', 'Puntaje']):
        hdr_cell(tbl.rows[0].cells[i], h)

    # Filas de datos
    row_idx = 1
    for sec in secciones:
        for j, q in enumerate(sec['preguntas']):
            dr = tbl.rows[row_idx]

            c0 = dr.cells[0]
            add_vmerge(c0, restart=(j == 0))
            if j == 0:
                p = c0.paragraphs[0]; nospace(p, 2, 2)
                run(p, sec['titulo_corto'], size=9)
                c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                for r in c0.paragraphs[0].runs: r.text = ''

            for col_idx, (val, bold) in enumerate([
                (q.get('habilidad',''), False),
                (str(row_idx), False),
                (q.get('clave',''), True),
                ('1', False),
            ], start=1):
                c = dr.cells[col_idx]
                p = c.paragraphs[0]; nospace(p, 2, 2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run(p, val, bold=bold, size=9)

            row_idx += 1

    # Fila total
    tr = tbl.rows[row_idx]
    tr.cells[0].merge(tr.cells[4])
    p = tr.cells[0].paragraphs[0]; nospace(p, 4, 4)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(p, f"Total:  {total_q} puntos", bold=True, size=10)

    set_cell_widths(tbl, COLS)


# ─── función principal ────────────────────────────────────────────────────────

def generar(
    asignatura    = "Matemática",
    profesor      = "Manuel Astudillo Figueroa",
    n_eval        = "5",
    titulo_eval   = "EVALUACIÓN N°5 PROBABILIDADES",
    curso         = "6° básico",
    anio          = "2026",
    exigencia     = "60%",
    puntaje_ideal = "16 pts.",
    oa_codigo     = "OA23",
    oa_texto      = "Conjeturar acerca de las tendencias de resultados obtenidos en repeticiones de un mismo experimento con dados, monedas u otros de manera manual. (adaptado)",
    secciones     = None,
    output_path   = None,
):
    secciones = secciones or SECCIONES
    doc = Document()
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(10)

    # Página: Legal, márgenes ajustados
    for sec in doc.sections:
        sec.page_width    = Emu(12240 * 635)
        sec.page_height   = Emu(20160 * 635)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(1.5)

    # ── Página 1: Evaluación ──────────────────────────────────────────────────
    add_school_header(doc, asignatura, curso, profesor)
    add_eval_title(doc, titulo_eval, asignatura, curso, profesor)
    add_info_table(doc, curso, anio, exigencia, puntaje_ideal)
    add_oa_line(doc, oa_codigo, oa_texto)
    add_general_instructions(doc, exigencia)

    n = 1
    for sec in secciones:
        add_section_header(doc, sec['romano'], sec['titulo'],
                           instruccion_seccion=sec.get('instruccion'))
        for q in sec['preguntas']:
            add_question(doc, n, q['enunciado'], q['opciones'])
            n += 1

    add_footer_line(doc)

    # ── Página 2: Pauta ───────────────────────────────────────────────────────
    doc.add_page_break()
    add_school_header(doc, asignatura, curso, profesor)
    add_pauta_title(doc, n_eval, titulo_eval, curso)
    add_pauta_oa_box(doc, oa_codigo, oa_texto)
    add_pauta_table(doc, secciones)
    add_footer_line(doc)

    if output_path is None:
        import os
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        output_path = os.path.join(base, f"{curso} {asignatura} E{n_eval}. {oa_codigo}.docx")
    doc.save(output_path)
    return output_path


# ─── datos (réplica del PDF) ──────────────────────────────────────────────────

SECCIONES = [
    {
        'romano': 'I.',
        'titulo': 'Describen un diagrama de árbol por medio de ejemplos',
        'titulo_corto': 'Describen un diagrama de árbol',
        'instruccion': 'Marca con una X la alternativa correcta.',
        'preguntas': [
            {'enunciado': '¿Qué representa el siguiente diagrama si C = cara y S = sello?',
             'opciones': ['El lanzamiento de una moneda dos veces.',
                          'El lanzamiento de tres monedas.',
                          'El lanzamiento de una moneda tres veces.',
                          'El lanzamiento de una moneda cinco veces.'],
             'clave': 'c', 'habilidad': 'Recordar'},
            {'enunciado': 'Si se lanza una moneda y un dado, ¿cuántas combinaciones posibles hay?',
             'opciones': ['6', '8', '10', '12'],
             'clave': 'd', 'habilidad': 'Comprender'},
            {'enunciado': '¿Qué se puede observar en un diagrama de árbol?',
             'opciones': ['Solo los resultados ganadores.',
                          'Los resultados posibles de un experimento.',
                          'Las respuestas incorrectas.',
                          'Los números que más se repiten.'],
             'clave': 'b', 'habilidad': 'Comprender'},
            {'enunciado': 'Al lanzar dos monedas, ¿cuántas combinaciones distintas se pueden obtener?',
             'opciones': ['2', '3', '4', '6'],
             'clave': 'c', 'habilidad': 'Aplicar'},
        ],
    },
    {
        'romano': 'II.',
        'titulo': 'Enumeran resultados posibles de lanzamientos de monedas o dados con ayuda de un diagrama de árbol.',
        'titulo_corto': 'Enumeran resultados posibles',
        'instruccion': 'Marca con una X la alternativa correcta.',
        'preguntas': [
            {'enunciado': 'Al lanzar dos veces un dado de seis caras, ¿en cuántas combinaciones se obtiene al menos un 1?',
             'opciones': ['2', '10', '11', '24'],
             'clave': 'c', 'habilidad': 'Aplicar'},
            {'enunciado': 'Si lanzas dos monedas, ¿cuál de las siguientes opciones representa los posibles resultados?',
             'opciones': ['(C, C), (C, S), (S, C), (S, S)',
                          '(1, 2), (2, 3), (3, 4), (4, 5)',
                          '(A, B), (B, C), (C, D)',
                          '(C, S), (S, C)'],
             'clave': 'a', 'habilidad': 'Recordar'},
            {'enunciado': 'Observa el siguiente experimento: Se lanzan dos monedas y se anota si sale cara (C) o sello (S). ¿Cuál de los siguientes conjuntos muestra todos los resultados posibles?',
             'opciones': ['(C), (S)',
                          '(C, C), (C, S), (S, C), (S, S)',
                          '(1, 2, 3, 4, 5, 6)',
                          '(C, S, C)'],
             'clave': 'b', 'habilidad': 'Aplicar'},
            {'enunciado': 'Observa este diagrama de árbol de una moneda y un dado. ¿Qué representa cada rama del diagrama?',
             'opciones': ['Una combinación posible de los lanzamientos.',
                          'Un número par.',
                          'Solo los sellos.',
                          'Solo los números impares.'],
             'clave': 'a', 'habilidad': 'Comprender'},
        ],
    },
    {
        'romano': 'III.',
        'titulo': 'Realizan de manera repetitiva experimentos con monedas para conjeturar acerca de las tendencias de los resultados.',
        'titulo_corto': 'Realizan experimentos con monedas',
        'instruccion': 'Marca con una X la alternativa correcta.',
        'preguntas': [
            {'enunciado': 'Si lanzas una moneda una vez, ¿cuál es la probabilidad de que salga sello?',
             'opciones': ['25%', '50%', '75%', '100%'],
             'clave': 'b', 'habilidad': 'Recordar'},
            {'enunciado': 'Si lanzas una moneda 10 veces y obtienes 6 caras y 4 sellos, ¿qué puedes concluir?',
             'opciones': ['Es más fácil obtener sello.',
                          'Es más fácil obtener cara.',
                          'Ambos resultados son posibles y cercanos al 50%.',
                          'Las monedas no tienen caras.'],
             'clave': 'c', 'habilidad': 'Analizar'},
            {'enunciado': 'Si lanzas una moneda 100 veces, ¿qué porcentaje aproximado se espera de caras?',
             'opciones': ['10%', '25%', '50%', '100%'],
             'clave': 'c', 'habilidad': 'Aplicar'},
            {'enunciado': 'Si lanzas una moneda 4 veces, ¿en cuántas ocasiones podrías esperar que salga sello?',
             'opciones': ['1', '2', '3', '4'],
             'clave': 'b', 'habilidad': 'Comprender'},
        ],
    },
    {
        'romano': 'IV.',
        'titulo': 'Conjeturan acerca de porcentajes de ocurrencia de eventos relativos a lanzamientos de monedas o dados.',
        'titulo_corto': 'Conjeturan porcentajes de ocurrencia',
        'instruccion': 'Marca con una X la alternativa correcta.',
        'preguntas': [
            {'enunciado': 'Si en un dado los números pares son 2, 4 y 6, ¿qué fracción representa sacar un número par?',
             'opciones': ['1/6', '2/6', '3/6', '6/6'],
             'clave': 'c', 'habilidad': 'Aplicar'},
            {'enunciado': '¿Qué porcentaje representa esa fracción (3/6)?',
             'opciones': ['25%', '33%', '50%', '75%'],
             'clave': 'c', 'habilidad': 'Aplicar'},
            {'enunciado': 'Si lanzas un dado 100 veces, ¿en cuántas veces aproximadamente esperas obtener un número impar?',
             'opciones': ['10 veces', '25 veces', '50 veces', '100 veces'],
             'clave': 'c', 'habilidad': 'Comprender'},
            {'enunciado': 'Si lanzas una moneda 20 veces, ¿cuántas veces aproximadamente obtendrás sello?',
             'opciones': ['5 veces', '10 veces', '15 veces', '20 veces'],
             'clave': 'b', 'habilidad': 'Aplicar'},
        ],
    },
]

if __name__ == '__main__':
    path = generar()
    print(f"Guardado: {path}")
